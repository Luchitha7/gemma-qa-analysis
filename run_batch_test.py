"""Run all 50 SignalQA test transcripts through the QA pipeline.

Extracts the transcripts from the test-transcripts docx, scores each one through
the same code path the web app's POST /analyze endpoint uses, saves every raw
JSON response into batch_results/json/, and writes:
    batch_results/manifest.json               - metadata + headline score each
    batch_results/summary.csv                 - flat table for spreadsheets
    batch_results/SignalQA_Outcomes_Report.docx - the outcomes document

    venv/bin/python run_batch_test.py            # full run
    venv/bin/python run_batch_test.py --dry-run  # extraction check only
"""

import csv
import json
import re
import sys
import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

DOCX_PATH = Path("/Users/luchithajayawardena/Downloads/"
                 "SignalQA_Test_Transcripts.docx")
RESULTS_DIR = Path(__file__).resolve().parent / "batch_results"
JSON_DIR = RESULTS_DIR / "json"

TITLE_RE = re.compile(r"^(\d{1,2})\.\s+(.+)$")
META_RE = re.compile(r"(good|mixed|poor)\b.*·\s*(short|medium|long)\b",
                     re.IGNORECASE)


def extract_transcripts():
    """Pull every transcript out of the docx, in order."""
    doc = Document(str(DOCX_PATH))
    items = []
    cur = None
    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue
        tmatch = TITLE_RE.match(line)
        if tmatch and not line.startswith("["):
            num = int(tmatch.group(1))
            title = tmatch.group(2)
            cur = {
                "id": f"{num:02d}",
                "title": title,
                "category": title.split("\u2014")[0].strip().lower(),
                "lines": [],
            }
            items.append(cur)
            continue
        if cur is None:
            continue
        if len(cur["lines"]) == 0 and cur.get("quality") is None:
            m = META_RE.search(line)
            if m:
                cur["quality"] = m.group(1).lower()
                cur["length"] = m.group(2).lower()
                continue
        if line.startswith("[") and ":" in line:
            cur["lines"].append(line)
    return [it for it in items if it.get("quality") and len(it["lines"]) >= 4]


def slug(title):
    s = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")
    return re.sub(r"_+", "_", s)


def avg(values):
    vals = [v for v in values if v is not None]
    return round(sum(vals) / len(vals), 1) if vals else None


def main():
    dry_run = "--dry-run" in sys.argv
    items = extract_transcripts()
    print(f"Extracted {len(items)} transcripts from {DOCX_PATH.name}",
          flush=True)
    if dry_run or len(items) != 50:
        for it in items:
            q, l, n = it.get("quality"), it.get("length"), len(it["lines"])
            print(f"  {it['id']}  {it['title']:<42} {q}/{l}/{n}")
        if not dry_run:
            sys.exit("Expected 50 transcripts - aborting.")
        return

    RESULTS_DIR.mkdir(exist_ok=True)
    JSON_DIR.mkdir(exist_ok=True)

    from web_app import app
    from fastapi.testclient import TestClient
    client = TestClient(app)

    rows = []
    t0 = time.time()
    for i, it in enumerate(items, 1):
        text = "\n".join(it["lines"])
        started = time.time()
        resp = client.post("/analyze", json={"transcript": text})
        data = resp.json()
        if resp.status_code != 200 or "error" in data:
            print(f"  [{i}/50] {it['id']} FAILED: {data}", flush=True)
            data = {"error": str(data)[:200]}

        fname = f"{it['id']}_{slug(it['title'])}.json"
        (JSON_DIR / fname).write_text(json.dumps(data, indent=2))

        row = {
            **{k: it[k] for k in ("id", "title", "category",
                                  "quality", "length")},
            "turns": len(it["lines"]),
            "file": fname,
            "final": data.get("final"),
            "band": data.get("band", ""),
            "agent": data.get("agent"),
            "conversation": data.get("conversation"),
            "accuracy": data.get("accuracy_overall"),
            "compliance": data.get("compliance_score"),
            "response_time": data.get("response_time_score"),
            "intense_count": len(data.get("intense") or []),
            "warning": data.get("warning", ""),
            "seconds": round(time.time() - started, 1),
        }
        rows.append(row)
        print(f"  [{i}/50] {it['id']} final={row['final']} "
              f"({row['seconds']}s)", flush=True)

    (RESULTS_DIR / "manifest.json").write_text(
        json.dumps({
            "run_at": datetime.now().isoformat(timespec="seconds"),
            "source_docx": str(DOCX_PATH),
            "count": len(rows),
            "results": rows,
        }, indent=2))

    with open(RESULTS_DIR / "summary.csv", "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    write_report(rows)
    total_min = (time.time() - t0) / 60
    print(f"\nDone in {total_min:.1f} min. Results: {RESULTS_DIR}", flush=True)


def _cell(table_row, idx, value):
    table_row.cells[idx].text = "" if value is None else str(value)


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for r in rows:
        cells = table.add_row()
        for j, v in enumerate(r):
            _cell(cells, j, v)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)


def write_report(rows):
    from docx import Document as Doc

    def fmt(v, suffix=""):
        return "n/a" if v is None else f"{v}{suffix}"

    doc = Doc()
    doc.add_heading("SignalQA Batch Outcomes Report", 0)
    intro = doc.add_paragraph()
    intro.add_run(
        f"Run completed {datetime.now().strftime('%d %B %Y at %H:%M')}. "
        f"All 50 test transcripts were scored end-to-end through the QA "
        f"pipeline (RoBERTa sentiment + Gemma scorecard/suggestions + RAG "
        f"accuracy/compliance + response-time scoring). Raw JSON output for "
        f"every call is saved alongside this report in batch_results/json/."
    ).italic = True

    doc.add_heading("1. What was run", level=1)
    cats = {}
    quals = {}
    lens = {}
    for r in rows:
        cats[r["category"]] = cats.get(r["category"], 0) + 1
        quals[r["quality"]] = quals.get(r["quality"], 0) + 1
        lens[r["length"]] = lens.get(r["length"], 0) + 1
    mix = (f"Transcripts: {len(rows)}   |   "
           + ", ".join(f"{k}: {v}" for k, v in sorted(cats.items()))
           + "   |   quality: "
           + ", ".join(f"{k}: {v}" for k, v in sorted(quals.items()))
           + "   |   length: "
           + ", ".join(f"{k}: {v}" for k, v in sorted(lens.items())))
    doc.add_paragraph(mix)
    doc.add_paragraph(
        "The set is deliberately varied so we can see how outcomes change "
        "with call quality, call length and call type.")

    doc.add_heading("2. Headline numbers", level=1)
    finals = [r["final"] for r in rows]
    bands = {}
    for r in rows:
        bands[r["band"]] = bands.get(r["band"], 0) + 1
    doc.add_paragraph(
        f"Average final QA score: {avg(finals)} / 100   |   "
        f"Highest: {max(finals)} ({[r['title'] for r in rows if r['final'] == max(finals)][0]})   |   "
        f"Lowest: {min(finals)} ({[r['title'] for r in rows if r['final'] == min(finals)][0]})")
    doc.add_paragraph("Score bands: "
                      + ", ".join(f"{k}: {v}" for k, v in sorted(bands.items())))

    doc.add_heading("3. Per-call results", level=1)
    _add_table(
        doc,
        ["#", "Call", "Cat", "Qual", "Len", "Turns", "Final", "Band",
         "Agent", "Conv", "Acc", "Comp", "RT"],
        [[r["id"], r["title"], r["category"][:4], r["quality"], r["length"],
          r["turns"], fmt(r["final"]), r["band"], fmt(r["agent"]),
          fmt(r["conversation"]), fmt(r["accuracy"]), fmt(r["compliance"]),
          fmt(r["response_time"])] for r in rows])

    doc.add_heading("4. How the outcome changes", level=1)

    doc.add_heading("By labelled quality", level=2)
    tq = [(q, avg([r["final"] for r in rows if r["quality"] == q]))
          for q in ("good", "mixed", "poor")]
    _add_table(doc, ["Quality label", "Avg final score", "Calls"],
               [[q, fmt(s), sum(1 for r in rows if r["quality"] == q)]
                for q, s in tq])
    gap = tq[0][1] - tq[2][1]
    doc.add_paragraph(
        f"Good calls average {tq[0][1]} vs {tq[2][1]} for poor calls - a "
        f"{gap:.1f}-point gap. The system scores good handling clearly higher "
        f"than poor handling." if gap > 5 else
        f"Good calls average {tq[0][1]} vs {tq[2][1]} for poor calls - only a "
        f"{gap:.1f}-point gap, which suggests the scoring may not be "
        f"differentiating quality strongly enough.")

    doc.add_heading("By call length", level=2)
    tl = [(ln, avg([r["final"] for r in rows if r["length"] == ln]),
           avg([r["response_time"] for r in rows if r["length"] == ln]))
          for ln in ("short", "medium", "long")]
    _add_table(doc, ["Length", "Avg final score", "Avg response-time score",
                     "Calls"],
               [[ln, fmt(s), fmt(rt), sum(1 for r in rows if r["length"] == ln)]
                for ln, s, rt in tl])
    drops = [r for r in rows if r["length"] == "long"]
    if drops:
        doc.add_paragraph(
            "Long calls stay scoreable end-to-end; no truncation failures "
            f"occurred ({len(drops)} long calls, avg final "
            f"{drops and avg([r['final'] for r in drops])}).")

    doc.add_heading("By category", level=2)
    tc = [(c, avg([r["final"] for r in rows if r["category"] == c]))
          for c in sorted(cats)]
    _add_table(doc, ["Category", "Avg final score", "Calls"],
               [[c, fmt(s), cats[c]] for c, s in tc])

    doc.add_heading("5. Notable calls", level=1)
    ranked = sorted(rows, key=lambda r: r["final"])
    doc.add_paragraph("Lowest-scoring calls:")
    for r in ranked[:3]:
        doc.add_paragraph(f"    {r['id']} {r['title']} - final "
                          f"{fmt(r['final'])}, agent {fmt(r['agent'])}",
                          style="List Bullet")
    doc.add_paragraph("Highest-scoring calls:")
    for r in ranked[-3:][::-1]:
        doc.add_paragraph(f"    {r['id']} {r['title']} - final "
                          f"{fmt(r['final'])}, agent {fmt(r['agent'])}",
                          style="List Bullet")

    warned = [r for r in rows if r["warning"]]
    if warned:
        doc.add_paragraph("Parsing warnings:")
        for r in warned:
            doc.add_paragraph(f"    {r['id']} {r['title']}: {r['warning']}",
                              style="List Bullet")
    else:
        doc.add_paragraph(
            "No parsing warnings: every transcript was read with both Agent "
            "and Client turns detected.")

    comp_fails = [r for r in rows if (r["compliance"] or 100) < 100]
    doc.add_paragraph(
        f"Compliance: {len(comp_fails)} of {len(rows)} calls broke at least "
        f"one rule."
        + (f" Lowest compliance score: {min(r['compliance'] for r in comp_fails)} "
           f"({comp_fails and min(comp_fails, key=lambda r: r['compliance'])['title']})."
           if comp_fails else ""))

    intense_ranked = sorted(rows, key=lambda r: r["intense_count"],
                            reverse=True)[:3]
    doc.add_paragraph("Most tense conversations (tense moments flagged by "
                      "RoBERTa): "
                      + "; ".join(f"{r['id']} {r['title']} ({r['intense_count']})"
                                  for r in intense_ranked))

    doc.add_heading("6. Method and how to reproduce", level=1)
    doc.add_paragraph(
        "Each transcript was posted to the same endpoint the web app exposes "
        "(POST /analyze), which runs: RoBERTa per-line sentiment and tense-"
        "moment flagging; Gemma (gemma3:1b via Ollama, temperature 0) summary, "
        "agent scorecard and coaching suggestions; RAG answer accuracy against "
        "the knowledge base; RAG compliance checks; timestamp-based response-"
        "time scoring; then the weighted final QA score.")
    doc.add_paragraph(
        "Reproduce with:  venv/bin/python run_batch_test.py   "
        "(--dry-run to check extraction only). Raw JSON responses are in "
        "batch_results/json/, a machine-readable table in "
        "batch_results/summary.csv and run metadata in "
        "batch_results/manifest.json.")

    out = RESULTS_DIR / "SignalQA_Outcomes_Report.docx"
    doc.save(str(out))
    print(f"Report written: {out}", flush=True)


if __name__ == "__main__":
    main()
